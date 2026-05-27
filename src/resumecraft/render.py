from dataclasses import dataclass
from re import Pattern

from docx.document import Document as DocumentObject
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Length, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from resumecraft.styles import (
    BODY_SIZE,
    COMPANY_SIZE,
    PAGE_WIDTH,
    SECTION_HEADING_SIZE,
    TECH_LINE_SIZE,
    ResolvedStyle,
)
from resumecraft.utils import add_bottom_border, add_hyperlink, keep_with_next


@dataclass
class RenderContext:
    doc: DocumentObject
    style: ResolvedStyle
    bold_pattern: Pattern[str] | None
    bold_keywords: frozenset[str]
    headings: dict[str, str]
    ats: bool


def heading_for(ctx: RenderContext, section: str, default: str) -> str:
    return ctx.headings.get(section, default)


def run(
    ctx: RenderContext,
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: Length | None = None,
    font: str | None = None,
) -> Run:
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = size or BODY_SIZE
    r.font.name = font or ctx.style.font_name
    return r


def add_section_heading(ctx: RenderContext, text: str) -> Paragraph:
    p = ctx.doc.add_paragraph()
    p.paragraph_format.space_before = ctx.style.section_space_before
    p.paragraph_format.space_after = ctx.style.section_space_after
    r = p.add_run(text)
    r.bold = True
    r.font.size = SECTION_HEADING_SIZE
    r.font.name = ctx.style.font_name
    r.font.color.rgb = ctx.style.heading_color
    if not ctx.ats:
        add_bottom_border(p)
    keep_with_next(p)
    return p


def add_two_column_line(
    ctx: RenderContext,
    left: str,
    right: str,
    *,
    left_bold: bool = True,
    left_italic: bool = False,
    left_size: Length | None = None,
) -> Paragraph:
    p = ctx.doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)

    run(ctx, p, left, bold=left_bold, italic=left_italic,
        size=left_size or COMPANY_SIZE)
    if ctx.ats:
        run(ctx, p, "  -  ", size=BODY_SIZE)
    else:
        p.paragraph_format.tab_stops.add_tab_stop(
            PAGE_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
    run(ctx, p, right, size=BODY_SIZE)

    keep_with_next(p)
    return p


def add_rich_bullet(ctx: RenderContext, text: str) -> Paragraph:
    p = ctx.doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = ctx.style.bullet_space
    p.paragraph_format.space_before = ctx.style.bullet_space

    if ctx.bold_pattern:
        for part in ctx.bold_pattern.split(text):
            if not part:
                continue
            r = p.add_run(part)
            r.font.size = BODY_SIZE
            r.font.name = ctx.style.font_name
            if part in ctx.bold_keywords:
                r.bold = True
    else:
        run(ctx, p, text)

    return p


def add_project_header(ctx: RenderContext, name: str, subtitle: str) -> Paragraph:
    p = ctx.doc.add_paragraph()
    p.paragraph_format.space_before = ctx.style.job_space_before
    p.paragraph_format.space_after = Pt(2)
    run(ctx, p, name, bold=True, size=COMPANY_SIZE)
    run(ctx, p, f"    {subtitle}", size=BODY_SIZE)
    keep_with_next(p)
    return p


def add_tech_line(ctx: RenderContext, text: str) -> None:
    p = ctx.doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    label = f"Tech: {text}" if ctx.ats else text
    r = p.add_run(label)
    r.italic = not ctx.ats
    r.font.size = TECH_LINE_SIZE
    r.font.name = ctx.style.font_name
    if not ctx.ats:
        r.font.color.rgb = ctx.style.tech_line_color
    keep_with_next(p)


def add_link_line(ctx: RenderContext, label: str, url: str) -> None:
    p = ctx.doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(ctx, p, label, size=TECH_LINE_SIZE)
    add_hyperlink(p, url, url, ctx.style.link_color, ctx.style.font_name, TECH_LINE_SIZE)
    keep_with_next(p)
