"""Section classes that render parts of a resume to docx.

Each section knows its default heading, whether it should be skipped
(empty), and how to render itself given a RenderContext.
"""

from abc import ABC, abstractmethod

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from resumecraft.models import Project, Resume
from resumecraft.render import (
    RenderContext,
    add_link_line,
    add_project_header,
    add_rich_bullet,
    add_section_heading,
    add_tech_line,
    add_two_column_line,
    heading_for,
    run,
)
from resumecraft.styles import (
    BODY_SIZE,
    CONTACT_SIZE,
    NAME_SIZE,
    PAGE_WIDTH,
    PHOTO_COL_WIDTH,
    PHOTO_MAX_HEIGHT,
    PHOTO_MAX_WIDTH,
)
from resumecraft.utils import add_hyperlink, load_photo_image, remove_table_borders


class Section(ABC):
    name: str
    default_heading: str

    def heading(self, ctx: RenderContext) -> str:
        return heading_for(ctx, self.name, self.default_heading)

    def is_empty(self, resume: Resume) -> bool:
        return False

    @abstractmethod
    def render(self, ctx: RenderContext, resume: Resume) -> None:
        ...


class HeaderSection(Section):
    """Name + contact line. Always rendered, not part of section_order."""

    name = "header"
    default_heading = ""

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        if resume.photo and not ctx.ats:
            self._render_with_photo(ctx, resume)
        else:
            self._render_centered(ctx, resume)

    def _render_centered(self, ctx: RenderContext, resume: Resume) -> None:
        p = ctx.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run(ctx, p, resume.name, bold=True, size=NAME_SIZE)

        self._add_contact_line(ctx, resume, WD_ALIGN_PARAGRAPH.CENTER)

    def _render_with_photo(self, ctx: RenderContext, resume: Resume) -> None:
        image_stream = load_photo_image(resume.photo)  # type: ignore[arg-type]

        table = ctx.doc.add_table(rows=1, cols=2)
        remove_table_borders(table)
        table.columns[0].width = PHOTO_COL_WIDTH
        table.columns[1].width = PAGE_WIDTH - PHOTO_COL_WIDTH

        photo_cell = table.cell(0, 0)
        photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        photo_p = photo_cell.paragraphs[0]
        photo_p.paragraph_format.space_before = Pt(0)
        photo_p.paragraph_format.space_after = Pt(0)
        # Pick the constraining dimension so the image fits within the passport-sized box
        # without distortion, regardless of its native aspect ratio
        from docx.image.image import Image as DocxImage

        blob = image_stream.read()
        image_stream.seek(0)
        img = DocxImage.from_blob(blob)
        if img.px_width / img.px_height >= PHOTO_MAX_WIDTH / PHOTO_MAX_HEIGHT:
            photo_p.add_run().add_picture(image_stream, width=PHOTO_MAX_WIDTH)
        else:
            photo_p.add_run().add_picture(image_stream, height=PHOTO_MAX_HEIGHT)

        info_cell = table.cell(0, 1)
        info_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        name_p = info_cell.paragraphs[0]
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_p.paragraph_format.space_after = Pt(2)
        name_p.paragraph_format.space_before = Pt(0)
        run(ctx, name_p, resume.name, bold=True, size=NAME_SIZE)

        self._add_contact_line(ctx, resume, WD_ALIGN_PARAGRAPH.LEFT, cell=info_cell)

    def _add_contact_line(
        self,
        ctx: RenderContext,
        resume: Resume,
        alignment: int,
        cell: object | None = None,
    ) -> None:
        p = cell.add_paragraph() if cell is not None else ctx.doc.add_paragraph()  # type: ignore[attr-defined]
        p.alignment = alignment
        p.paragraph_format.space_after = Pt(2)

        contact = resume.contact
        run(ctx, p, f"{contact.location}  |  ", size=CONTACT_SIZE)
        add_hyperlink(
            p, contact.email, f"mailto:{contact.email}",
            ctx.style.link_color, ctx.style.font_name, CONTACT_SIZE,
        )
        run(ctx, p, f"  |  {contact.phone}", size=CONTACT_SIZE)

        for link in contact.links:
            run(ctx, p, "  |  ", size=CONTACT_SIZE)
            add_hyperlink(p, link.label, link.url, ctx.style.link_color, ctx.style.font_name, CONTACT_SIZE)


class SummarySection(Section):
    name = "summary"
    default_heading = "PROFESSIONAL SUMMARY"

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        add_section_heading(ctx, self.heading(ctx))
        p = ctx.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run(ctx, p, resume.summary)


class ExperienceSection(Section):
    name = "experience"
    default_heading = "WORK EXPERIENCE"

    def is_empty(self, resume: Resume) -> bool:
        return not resume.experience

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        add_section_heading(ctx, self.heading(ctx))
        for i, exp in enumerate(resume.experience):
            p = add_two_column_line(ctx, exp.company, exp.location)
            p.paragraph_format.space_before = ctx.style.job_space_before if i == 0 else Pt(8)
            add_two_column_line(
                ctx, exp.title, exp.duration,
                left_bold=False, left_italic=True, left_size=BODY_SIZE,
            )
            for bullet in exp.bullets:
                add_rich_bullet(ctx, bullet)


class ProjectsSection(Section):
    """One class, three registered instances (projects / professional / personal)."""

    def __init__(self, attr: str, name: str, default_heading: str) -> None:
        self.attr = attr
        self.name = name
        self.default_heading = default_heading

    def is_empty(self, resume: Resume) -> bool:
        return not getattr(resume, self.attr)

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        projects: list[Project] = getattr(resume, self.attr)
        add_section_heading(ctx, self.heading(ctx))
        for proj in projects:
            add_project_header(ctx, proj.name, proj.subtitle)
            if proj.tech_stack:
                add_tech_line(ctx, proj.tech_stack)
            for link in proj.all_links:
                add_link_line(ctx, f"{link.label}: ", link.url)
            for bullet in proj.bullets:
                add_rich_bullet(ctx, bullet)


class SkillsSection(Section):
    name = "skills"
    default_heading = "SKILLS"

    def is_empty(self, resume: Resume) -> bool:
        return not resume.skills

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        add_section_heading(ctx, self.heading(ctx))
        for skill in resume.skills:
            p = ctx.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            run(ctx, p, f"{skill.category}: ", bold=True)
            run(ctx, p, skill.items)


class EducationSection(Section):
    name = "education"
    default_heading = "EDUCATION"

    def is_empty(self, resume: Resume) -> bool:
        return not resume.education

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        add_section_heading(ctx, self.heading(ctx))
        for edu in resume.education:
            p = add_two_column_line(ctx, edu.institution, edu.duration)
            p.paragraph_format.space_before = Pt(2)
            p2 = ctx.doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            run(ctx, p2, edu.degree, italic=True)


class CertificationsSection(Section):
    name = "certifications"
    default_heading = "CERTIFICATIONS"

    def is_empty(self, resume: Resume) -> bool:
        return not resume.certifications

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        add_section_heading(ctx, self.heading(ctx))
        for cert in resume.certifications:
            p = add_two_column_line(ctx, cert.name, cert.date)
            p.paragraph_format.space_before = Pt(2)
            p2 = ctx.doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            run(ctx, p2, cert.issuer, italic=True)
            for link in cert.all_links:
                add_link_line(ctx, f"{link.label}: ", link.url)


class AwardsSection(Section):
    name = "awards"
    default_heading = "AWARDS & ACHIEVEMENTS"

    def is_empty(self, resume: Resume) -> bool:
        return not resume.awards

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        add_section_heading(ctx, self.heading(ctx))
        for award in resume.awards:
            p = add_two_column_line(ctx, award.title, award.date)
            p.paragraph_format.space_before = Pt(2)
            p2 = ctx.doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            run(ctx, p2, award.issuer, italic=True)
            if award.description:
                p3 = ctx.doc.add_paragraph()
                p3.paragraph_format.space_before = Pt(0)
                run(ctx, p3, award.description)


class LanguagesSection(Section):
    name = "languages"
    default_heading = "LANGUAGES"

    def is_empty(self, resume: Resume) -> bool:
        return not resume.languages

    def render(self, ctx: RenderContext, resume: Resume) -> None:
        add_section_heading(ctx, self.heading(ctx))
        p = ctx.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run(ctx, p, resume.languages)


SECTION_REGISTRY: dict[str, Section] = {
    "summary": SummarySection(),
    "experience": ExperienceSection(),
    "projects": ProjectsSection("projects", "projects", "PROJECTS"),
    "professional_projects": ProjectsSection(
        "professional_projects", "professional_projects", "PROFESSIONAL PROJECTS",
    ),
    "personal_projects": ProjectsSection(
        "personal_projects", "personal_projects", "PERSONAL & OPEN SOURCE PROJECTS",
    ),
    "skills": SkillsSection(),
    "education": EducationSection(),
    "certifications": CertificationsSection(),
    "awards": AwardsSection(),
    "languages": LanguagesSection(),
}

DEFAULT_HEADINGS: dict[str, str] = {
    name: section.default_heading for name, section in SECTION_REGISTRY.items()
}
