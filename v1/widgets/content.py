"""Content widgets - the rendering primitives."""

from __future__ import annotations

from dataclasses import dataclass, field

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from .base import BuildContext, Widget
from .helpers import add_bottom_border, add_hyperlink, build_bold_pattern, load_image

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _apply_font(
    run: object,
    ctx: BuildContext,
    size: float | None,
    bold: bool,
    italic: bool,
    color: str | None,
) -> None:
    """Set font properties on a run. Uses ctx defaults when widget doesn't override."""
    font = run.font  # type: ignore[attr-defined]
    font.name = ctx.font_name
    font.size = Pt(size or ctx.font_size)
    font.bold = True if bold else None  # None = inherit from style
    font.italic = True if italic else None
    c = color or ctx.font_color
    if c:
        font.color.rgb = RGBColor.from_string(c.lstrip("#"))


@dataclass
class Text(Widget):
    """Single paragraph of text with optional formatting."""

    content: str
    size: float | None = None
    bold: bool = False
    italic: bool = False
    color: str | None = None
    align: str = "left"
    space_before: float | None = None
    space_after: float | None = None

    def build(self, ctx: BuildContext) -> None:
        p = ctx.container.add_paragraph()
        p.alignment = _ALIGN.get(self.align, WD_ALIGN_PARAGRAPH.LEFT)
        fmt = p.paragraph_format
        if self.space_before is not None:
            fmt.space_before = Pt(self.space_before)
        if self.space_after is not None:
            fmt.space_after = Pt(self.space_after)
        run = p.add_run(self.content)
        _apply_font(run, ctx, self.size, self.bold, self.italic, self.color)


@dataclass
class Heading(Widget):
    """Section heading - uppercase with bottom border by default."""

    content: str
    size: float = 13
    bold: bool = True
    uppercase: bool = True
    border_bottom: bool = True
    color: str | None = None
    align: str = "left"
    space_before: float = 6
    space_after: float = 2

    def build(self, ctx: BuildContext) -> None:
        p = ctx.container.add_paragraph()
        p.alignment = _ALIGN.get(self.align, WD_ALIGN_PARAGRAPH.LEFT)
        fmt = p.paragraph_format
        fmt.space_before = Pt(self.space_before)
        fmt.space_after = Pt(self.space_after)

        text = self.content.upper() if self.uppercase else self.content
        run = p.add_run(text)
        _apply_font(run, ctx, self.size, self.bold, False, self.color)

        if self.border_bottom:
            border_color = (self.color or ctx.font_color or "000000").lstrip("#")
            add_bottom_border(p, border_color)


@dataclass
class Bullet(Widget):
    """Bullet point with optional bold keyword highlighting."""

    content: str
    size: float | None = None
    bold_words: list[str] = field(default_factory=list)
    space_before: float = 0
    space_after: float = 0

    def build(self, ctx: BuildContext) -> None:
        p = ctx.container.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = Pt(self.space_before)
        fmt.space_after = Pt(self.space_after)

        bullet_run = p.add_run("• ")
        _apply_font(bullet_run, ctx, self.size, False, False, None)

        pattern = build_bold_pattern(self.bold_words) if self.bold_words else ctx.bold_pattern
        if pattern:
            parts = pattern.split(self.content)
            bold_set = {m.lower() for m in pattern.findall(self.content)}
            for part in parts:
                if not part:
                    continue
                is_bold = part.lower() in bold_set
                run = p.add_run(part)
                _apply_font(run, ctx, self.size, is_bold, False, None)
        else:
            run = p.add_run(self.content)
            _apply_font(run, ctx, self.size, False, False, None)


@dataclass
class TwoCol(Widget):
    """Left-right text on one line using a right-aligned tab stop."""

    left: str
    right: str
    size: float | None = None
    left_bold: bool = False
    left_italic: bool = False
    right_bold: bool = False
    right_italic: bool = False
    color: str | None = None
    space_before: float | None = None
    space_after: float | None = None

    def build(self, ctx: BuildContext) -> None:
        p = ctx.container.add_paragraph()
        fmt = p.paragraph_format
        if self.space_before is not None:
            fmt.space_before = Pt(self.space_before)
        if self.space_after is not None:
            fmt.space_after = Pt(self.space_after)

        page_width = (
            ctx.doc.sections[0].page_width
            - ctx.doc.sections[0].left_margin
            - ctx.doc.sections[0].right_margin
        )
        fmt.tab_stops.add_tab_stop(page_width, WD_TAB_ALIGNMENT.RIGHT)

        left_run = p.add_run(self.left)
        _apply_font(left_run, ctx, self.size, self.left_bold, self.left_italic, self.color)

        p.add_run("\t")

        right_run = p.add_run(self.right)
        _apply_font(right_run, ctx, self.size, self.right_bold, self.right_italic, self.color)


@dataclass
class Link(Widget):
    """Clickable hyperlink rendered as blue underlined text."""

    label: str
    url: str
    size: float | None = None
    color: str | None = None
    align: str = "left"

    def build(self, ctx: BuildContext) -> None:
        p = ctx.container.add_paragraph()
        p.alignment = _ALIGN.get(self.align, WD_ALIGN_PARAGRAPH.LEFT)
        add_hyperlink(
            p,
            self.label,
            self.url,
            font_name=ctx.font_name,
            font_size=self.size or ctx.font_size,
            link_color=(self.color or "0563C1").lstrip("#"),
        )


@dataclass
class Image(Widget):
    """Inline image from a local file path or URL. Dimensions in inches."""

    source: str
    width: float | None = None
    height: float | None = None
    align: str = "left"

    def build(self, ctx: BuildContext) -> None:
        stream = load_image(self.source)
        p = ctx.container.add_paragraph()
        p.alignment = _ALIGN.get(self.align, WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run()
        kwargs: dict[str, Inches] = {}
        if self.width:
            kwargs["width"] = Inches(self.width)
        if self.height:
            kwargs["height"] = Inches(self.height)
        run.add_picture(stream, **kwargs)


@dataclass
class Spacer(Widget):
    """Empty vertical space. Height in points."""

    height: float = 6

    def build(self, ctx: BuildContext) -> None:
        p = ctx.container.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        run = p.add_run()
        run.font.size = Pt(self.height)


@dataclass
class Divider(Widget):
    """Horizontal line spanning the container width."""

    color: str | None = None
    space_before: float = 2
    space_after: float = 2

    def build(self, ctx: BuildContext) -> None:
        p = ctx.container.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = Pt(self.space_before)
        fmt.space_after = Pt(self.space_after)
        border_color = (self.color or ctx.font_color or "000000").lstrip("#")
        add_bottom_border(p, border_color)
