"""Document - the root widget that creates and manages the docx file."""

from __future__ import annotations

import io
from dataclasses import dataclass, field

from docx import Document as new_docx_document
from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt

from ..templates.base import Template
from .base import BuildContext, Widget
from .helpers import build_bold_pattern


@dataclass
class Document:
    """Root of the widget tree. Creates a docx and renders all children."""

    children: list[Widget] = field(default_factory=list)
    template: Template = field(default_factory=Template)
    bold_words: list[str] = field(default_factory=list)

    def build(self) -> DocxDocument:
        """Create a python-docx Document and render all children into it."""
        tmpl = self.template
        doc = new_docx_document()

        for section in doc.sections:
            section.top_margin = Inches(tmpl.margin_top)
            section.bottom_margin = Inches(tmpl.margin_bottom)
            section.left_margin = Inches(tmpl.margin_left)
            section.right_margin = Inches(tmpl.margin_right)

        style = doc.styles["Normal"]
        style.font.name = tmpl.font
        style.font.size = Pt(tmpl.font_size)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        if doc.paragraphs:
            doc.paragraphs[0]._element.getparent().remove(
                doc.paragraphs[0]._element
            )

        ctx = BuildContext(
            doc=doc,
            container=doc,
            template=tmpl,
            font_name=tmpl.font,
            font_size=tmpl.font_size,
            bold_pattern=build_bold_pattern(self.bold_words),
        )

        for child in self.children:
            child.build(ctx)

        return doc

    def save(self, path: str) -> None:
        """Build and write to a .docx file."""
        self.build().save(path)

    def to_bytes(self) -> bytes:
        """Build and return raw docx bytes (useful for web responses)."""
        buf = io.BytesIO()
        self.build().save(buf)
        return buf.getvalue()
