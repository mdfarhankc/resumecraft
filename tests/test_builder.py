from pathlib import Path

from resumecraft.builder import DocxBuilder
from resumecraft.models import StyleOptions


class TestDocxBuilder:
    def test_build_returns_document(self, minimal_resume):
        builder = DocxBuilder(minimal_resume)
        doc = builder.build()
        assert doc is not None
        assert hasattr(doc, "paragraphs")

    def test_save_creates_file(self, minimal_resume, tmp_path):
        output = tmp_path / "resume.docx"
        builder = DocxBuilder(minimal_resume)
        result = builder.save(output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_save_creates_parent_dirs(self, minimal_resume, tmp_path):
        output = tmp_path / "nested" / "dir" / "resume.docx"
        builder = DocxBuilder(minimal_resume)
        builder.save(output)
        assert output.exists()

    def test_full_resume_builds(self, full_resume, tmp_path):
        output = tmp_path / "full.docx"
        builder = DocxBuilder(full_resume)
        builder.save(output)
        assert output.exists()
        assert output.stat().st_size > 0

    def test_document_has_content(self, full_resume):
        builder = DocxBuilder(full_resume)
        doc = builder.build()
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Smith" in text
        assert "WORK EXPERIENCE" in text
        assert "PROFESSIONAL PROJECTS" in text
        assert "PERSONAL & OPEN SOURCE PROJECTS" in text
        assert "SKILLS" in text
        assert "EDUCATION" in text
        assert "LANGUAGES" in text

    def test_minimal_resume_skips_empty_sections(self, minimal_resume):
        builder = DocxBuilder(minimal_resume)
        doc = builder.build()
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "John Doe" in text
        assert "PROFESSIONAL SUMMARY" in text
        assert "WORK EXPERIENCE" not in text
        assert "SKILLS" not in text
        assert "EDUCATION" not in text

    def test_bold_keywords_applied(self, full_resume):
        builder = DocxBuilder(full_resume)
        doc = builder.build()
        bold_runs = []
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and run.text.strip():
                    bold_runs.append(run.text.strip())
        # Section headings and keywords should be bold
        assert "FastAPI" in bold_runs
        assert "React" in bold_runs

    def test_page_margins(self, minimal_resume):
        builder = DocxBuilder(minimal_resume)
        builder.build()
        from resumecraft.styles import BOTTOM_MARGIN, LEFT_MARGIN, RIGHT_MARGIN, TOP_MARGIN
        for section in builder.doc.sections:
            assert section.top_margin == TOP_MARGIN
            assert section.bottom_margin == BOTTOM_MARGIN
            assert section.left_margin == LEFT_MARGIN
            assert section.right_margin == RIGHT_MARGIN

    def test_sample_json_builds(self, sample_json_path, tmp_path):
        from resumecraft.models import Resume
        resume = Resume.from_json(str(sample_json_path))
        output = tmp_path / "sample.docx"
        DocxBuilder(resume).save(output)
        assert output.exists()

    def test_custom_section_order(self, full_resume):
        full_resume.section_order = ["skills", "experience", "summary"]
        builder = DocxBuilder(full_resume)
        doc = builder.build()
        text = "\n".join(p.text for p in doc.paragraphs)
        skills_pos = text.index("SKILLS")
        exp_pos = text.index("WORK EXPERIENCE")
        summary_pos = text.index("PROFESSIONAL SUMMARY")
        assert skills_pos < exp_pos < summary_pos
        # Sections not in order should be absent
        assert "EDUCATION" not in text
        assert "LANGUAGES" not in text

    def test_default_section_order(self, full_resume):
        # When section_order is None, all sections render in default order
        full_resume.section_order = None
        builder = DocxBuilder(full_resume)
        doc = builder.build()
        text = "\n".join(p.text for p in doc.paragraphs)
        summary_pos = text.index("PROFESSIONAL SUMMARY")
        exp_pos = text.index("WORK EXPERIENCE")
        skills_pos = text.index("SKILLS")
        assert summary_pos < exp_pos < skills_pos

    def test_style_font(self, minimal_resume, tmp_path):
        minimal_resume.style = StyleOptions(font="georgia")
        builder = DocxBuilder(minimal_resume)
        builder.build()
        assert builder.doc.styles["Normal"].font.name == "Georgia"

    def test_style_color(self, minimal_resume):
        minimal_resume.style = StyleOptions(color="navy")
        builder = DocxBuilder(minimal_resume)
        doc = builder.build()
        # Check heading color is navy (0, 32, 96)
        from docx.shared import RGBColor
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text == "PROFESSIONAL SUMMARY":
                    assert run.font.color.rgb == RGBColor(0, 32, 96)

    def test_style_spacing_compact(self, full_resume, tmp_path):
        full_resume.style = StyleOptions(spacing="compact")
        output = tmp_path / "compact.docx"
        DocxBuilder(full_resume).save(output)
        assert output.exists()

    def test_style_spacing_relaxed(self, full_resume, tmp_path):
        full_resume.style = StyleOptions(spacing="relaxed")
        output = tmp_path / "relaxed.docx"
        DocxBuilder(full_resume).save(output)
        assert output.exists()

    def test_all_fonts_build(self, minimal_resume, tmp_path):
        for font in ["calibri", "arial", "times", "garamond", "georgia", "helvetica", "cambria"]:
            minimal_resume.style = StyleOptions(font=font)
            output = tmp_path / f"{font}.docx"
            DocxBuilder(minimal_resume).save(output)
            assert output.exists()

    def test_all_colors_build(self, minimal_resume, tmp_path):
        for color in ["black", "navy", "forest", "maroon", "slate", "royal"]:
            minimal_resume.style = StyleOptions(color=color)
            output = tmp_path / f"{color}.docx"
            DocxBuilder(minimal_resume).save(output)
            assert output.exists()
